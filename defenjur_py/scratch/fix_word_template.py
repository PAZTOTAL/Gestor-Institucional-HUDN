from docx import Document

def fix_template():
    doc = Document('defenjur_py/legal/templates_docx/ACCION_DE_TUTELA_modelo.docx')
    
    # 1. Reemplazos globales
    for p in doc.paragraphs:
        # Accionante
        if 'ANA LUISA ROSERO' in p.text:
            p.text = p.text.replace('ANA LUISA ROSERO', '{{ACCIONANTE}}')
            # Limpiar posibles residuos del nombre original si quedó algo
            if 'DÍAZ' in p.text: p.text = p.text.replace('DÍAZ', '')
            if 'DAZ' in p.text: p.text = p.text.replace('DAZ', '')
        
        # Juzgado
        if 'JUZGADO QUINTO DE FAMILIA' in p.text:
            p.text = p.text.replace('JUZGADO QUINTO DE FAMILIA DEL CIRCUITO PASTO - NARIÑO', '{{JUZGADO}}')
            p.text = p.text.replace('JUZGADO QUINTO DE FAMILIA DEL CIRCUITO PASTO - NARIO', '{{JUZGADO}}')
            # Fallback
            if 'JUZGADO QUINTO' in p.text: p.text = '{{JUZGADO}}'
            
        # Numero proceso
        if '2026-00140' in p.text:
            p.text = p.text.replace('2026-00140', '{{NUM_PROCESO}}')
            
        # Fecha
        if '27 de abril de 2026' in p.text:
            p.text = p.text.replace('27 de abril de 2026', '{{FECHA_HOY}}')

    # 2. Reemplazo de bloque de hechos
    has_hechos = False
    for p in doc.paragraphs:
        if 'FRENTE AL PRIMER HECHO' in p.text:
            p.text = '{{HECHOS}}'
            has_hechos = True
            continue
        
        if has_hechos:
            if 'PETICIONES' in p.text:
                has_hechos = False
            elif 'FRENTE AL' in p.text or 'Es cierto' in p.text or 'No me consta' in p.text or 'consideración' in p.text:
                p.text = ''

    doc.save('defenjur_py/legal/templates_docx/ACCION_DE_TUTELA_modelo.docx')
    print("Template fixed successfully")

if __name__ == '__main__':
    fix_template()
