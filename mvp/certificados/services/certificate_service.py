from io import BytesIO
from copy import deepcopy
import os
import re
import time
import unicodedata
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from lxml import etree

from .date_utils import get_spanish_expedition_date

BASE_DIR = Path(__file__).resolve().parents[2]
TEMPLATES_DOCX_DIR = BASE_DIR / "templates_docx"


def score_corruption(text):
    return len(re.findall(r"Ã|Â|â|�|[\u0080-\u009f]", text))


def repair_common_mojibake(text):
    replacements = {
        "Ã¡": "á",
        "Ã©": "é",
        "Ã­": "í",
        "Ã³": "ó",
        "Ãº": "ú",
        "Ã": "Á",
        "Ã‰": "É",
        "Ã": "Í",
        "Ã“": "Ó",
        "Ãš": "Ú",
        "Ã±": "ñ",
        "Ã‘": "Ñ",
        "Ã¼": "ü",
        "Ãœ": "Ü",
        "â€“": "–",
        "â€”": "—",
        "â€œ": "“",
        "â€": "”",
        "â€˜": "‘",
        "â€™": "’",
        "Â": "",
        "Ã?O": "ÑO",
        "Ã?o": "ño",
        "Ã?N": "ÓN",
        "Ã?n": "ón",
    }
    output = text
    for source, target in replacements.items():
        output = output.replace(source, target)
    return output


def repair_lost_accent_question_marks(text):
    replacements = {
        "GESTI?N": "GESTIÓN",
        "INFORMACI?N": "INFORMACIÓN",
        "ATENCI?N": "ATENCIÓN",
        "SECCI?N": "SECCIÓN",
        "OPERACI?N": "OPERACIÓN",
        "ADMINISTRACI?N": "ADMINISTRACIÓN",
        "CONTRATACI?N": "CONTRATACIÓN",
        "PRESTACI?N": "PRESTACIÓN",
        "RELACI?N": "RELACIÓN",
        "CERTIFICACI?N": "CERTIFICACIÓN",
        "T?CNICO": "TÉCNICO",
        "TECNOL?GICOS": "TECNOLÓGICOS",
        "NARI?O": "NARIÑO",
    }
    output = text
    for source, target in replacements.items():
        output = output.replace(source, target).replace(source.lower(), target.lower())
    return output


def sanitize_text(value):
    text = str(value or "")
    if not text:
        return text
    if re.search(r"Ã|Â|â|[\u0080-\u009f]", text):
        decoded = text.encode("latin1", errors="ignore").decode("utf-8", errors="ignore")
        if score_corruption(decoded) < score_corruption(text):
            text = decoded
    text = repair_common_mojibake(text)
    text = repair_lost_accent_question_marks(text)
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\uFFFD", "")
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text).strip()


def sanitize_template_data(value):
    if isinstance(value, str):
        return sanitize_text(value)
    if isinstance(value, list):
        return [sanitize_template_data(item) for item in value]
    if isinstance(value, dict):
        return {k: sanitize_template_data(v) for k, v in value.items()}
    return value


def template_path_by_gender(genero):
    normalized = str(genero or "").lower().strip()
    if normalized == "femenino":
        return TEMPLATES_DOCX_DIR / "plantilla_femenina.docx"
    return TEMPLATES_DOCX_DIR / "plantilla_masculina.docx"


def replace_tokens(text, mapping):
    output = text
    for key, value in mapping.items():
        output = output.replace(key, str(value))
    return output


def apply_mapping_to_paragraph(paragraph, mapping):
    original = paragraph.text
    replaced = replace_tokens(original, mapping)
    if replaced == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


def render_contracts_loop(document, contracts):
    loop_start = "{{#contratos}}"
    loop_end = "{{/contratos}}"

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            if loop_start not in row_text and loop_end not in row_text:
                continue

            template_tr = row._tr

            if not contracts:
                table._tbl.remove(template_tr)
                return

            for contract in contracts:
                contract_mapping = {
                    "{{#contratos}}": "",
                    "{{/contratos}}": "",
                    "{{contratoNo}}": contract.get("contratoNo", ""),
                    "{{firmaContrato}}": contract.get("firmaContrato", ""),
                    "{{fechaInicio}}": contract.get("fechaInicio", ""),
                    "{{fechaTerminacion}}": contract.get("fechaTerminacion", ""),
                    "{{valor}}": contract.get("valor", ""),
                }
                new_tr = deepcopy(template_tr)
                row_xml = etree.tostring(new_tr, encoding="unicode")
                row_xml = re.sub(r"<w:proofErr[^>]*/>", "", row_xml)

                while True:
                    collapsed = re.sub(
                        r"\{\{([^{}]*?)</w:t>.*?<w:t>([^{}]*?)\}\}",
                        r"{{\1\2}}",
                        row_xml,
                        flags=re.DOTALL,
                    )
                    if collapsed == row_xml:
                        break
                    row_xml = collapsed

                replacements_regex = {
                    r"\{\{\s*#\s*contratos\s*\}\}": "",
                    r"\{\{\s*/\s*contratos\s*\}\}": "",
                    r"\{\{\s*contratoNo\s*\}\}": contract.get("contratoNo", ""),
                    r"\{\{\s*firmaContrato\s*\}\}": contract.get("firmaContrato", ""),
                    r"\{\{\s*fechaInicio\s*\}\}": contract.get("fechaInicio", ""),
                    r"\{\{\s*fechaTerminacion\s*\}\}": contract.get("fechaTerminacion", ""),
                    r"\{\{\s*valor\s*\}\}": contract.get("valor", ""),
                }
                for pattern, value in replacements_regex.items():
                    row_xml = re.sub(pattern, xml_escape(str(value)), row_xml)

                new_tr_clean = etree.fromstring(row_xml.encode("utf-8"))
                template_tr.addprevious(new_tr_clean)
            table._tbl.remove(template_tr)
            return


def generate_certificate(data, genero):
    clean_data = sanitize_template_data(data)
    template_path = template_path_by_gender(genero)
    if not template_path.exists():
        raise ValueError(f"No existe la plantilla: {template_path}")

    doc = Document(str(template_path))
    fecha = get_spanish_expedition_date()
    context = {
        "nombre": clean_data.get("nombre"),
        "cedula": clean_data.get("cedula"),
        "cargo": clean_data.get("cargo"),
        "contratos": clean_data.get("contratos"),
        "objeto_ctto": clean_data.get("cargo"),
        "dia": fecha["dia"],
        "mes": fecha["mes"],
        "anio": fecha["anio"],
        "fecha_expedicion_texto": fecha["fecha_texto"],
    }

    try:
        scalar_mapping = {
            "{{nombre}}": context["nombre"] or "",
            "{{cedula}}": context["cedula"] or "",
            "{{cargo}}": context["cargo"] or "",
            "{{objeto_ctto}}": context["objeto_ctto"] or "",
            "{{dia}}": context["dia"],
            "{{mes}}": context["mes"],
            "{{anio}}": context["anio"],
            "{{fecha_expedicion_texto}}": context["fecha_expedicion_texto"] or "",
        }

        for paragraph in doc.paragraphs:
            apply_mapping_to_paragraph(paragraph, scalar_mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        apply_mapping_to_paragraph(paragraph, scalar_mapping)

        render_contracts_loop(doc, context["contratos"] or [])
    except Exception as error:
        raise ValueError(
            f"Error al renderizar la plantilla. Verifica etiquetas y bloque de contratos. Detalle: {error}"
        ) from error

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"certificado_{clean_data.get('cedula')}_{int(time.time() * 1000)}.docx"
    return output, filename
