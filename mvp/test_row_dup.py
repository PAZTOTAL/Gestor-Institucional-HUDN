import re
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from lxml import etree
import copy

def test():
    doc = Document('templates_docx/plantilla_masculina.docx')
    table = doc.tables[0]
    
    # Let's find the template row
    template_row = None
    for row in table.rows:
        text = " ".join(c.text for c in row.cells)
        if '{{#contratos}}' in text:
            template_row = row
            break
            
    print("Found template row:", template_row)
    
    mapping = {
        "{{#contratos}}": "",
        "{{/contratos}}": "",
        "{{contratoNo}}": "12345",
        "{{firmaContrato}}": "2024-01-01",
        "{{fechaInicio}}": "2024-01-02",
        "{{fechaTerminacion}}": "2024-12-31",
        "{{valor}}": "5000000"
    }
    
    # Duplicate row
    new_tr = copy.deepcopy(template_row._tr)
    from docx.table import _Row
    new_row = _Row(new_tr, table)
    
    from certificados.services.certificate_service import apply_mapping_to_paragraph
    
    for cell in new_row.cells:
        for paragraph in cell.paragraphs:
            apply_mapping_to_paragraph(paragraph, mapping)
            
    print("New Row Text:", " | ".join(c.text for c in new_row.cells))

test()
